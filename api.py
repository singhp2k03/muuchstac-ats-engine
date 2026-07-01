from fastapi import FastAPI, UploadFile, File, Form, HTTPException, Response
from fastapi.middleware.cors import CORSMiddleware
from contextlib import asynccontextmanager
from dataclasses import dataclass
from typing import List
import pypdf
import docx
import io
import json
import asyncio
import logging
import requests
import numpy as np
import hashlib 
from sklearn.metrics.pairwise import haversine_distances
from geopy.geocoders import Nominatim
from functools import lru_cache
from pydantic import BaseModel, Field
from google import genai
from google.genai import types

import os
from dotenv import load_dotenv
load_dotenv()

logging.basicConfig(level=logging.INFO, format="%(asctime)s | %(levelname)s | %(message)s")
logger = logging.getLogger(__name__)

# --- CONFIGURATION ---
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")   
# OPENAI_API_KEY = os.getenv("OPENAI_API_KEY") # Ready for future
# CUSTOM_API_URL = os.getenv("CUSTOM_API_URL") # Ready for future

GEMINI_MODEL = "gemini-3.1-flash-lite" 
MAX_CONCURRENT = 2        

client = genai.Client(api_key=GEMINI_API_KEY, http_options=types.HttpOptions(timeout=60 * 1000))
semaphore = asyncio.Semaphore(MAX_CONCURRENT)

# --- 1. THE HYBRID LOCATION ENGINE (FREE) ---
geolocator = Nominatim(user_agent="muuchstac_ats_hybrid")

@lru_cache(maxsize=1000)
def get_coordinates(city_name: str):
    if not city_name or city_name.lower() in ["not found", "n/a", "unknown"]:
        return None
    try:
        location = geolocator.geocode(f"{city_name}, India", timeout=5)
        if location:
            return (location.latitude, location.longitude, location.address)
    except Exception as e:
        logger.error(f"Geocoding failed for {city_name}: {e}")
    return None

def calculate_hybrid_location_score(candidate_city: str, target_city: str = "Borivali, Mumbai") -> dict:
    if not candidate_city or not target_city:
        return {"relevancy": "Unknown", "status": "Location unknown."}

    cand_clean = candidate_city.lower().replace("maharashtra", "").replace("india", "").strip(" ,.")
    target_clean = target_city.lower().replace("maharashtra", "").replace("india", "").strip(" ,.")
    
    cand_parts = [p.strip() for p in cand_clean.split(',') if p.strip()]
    target_parts = [p.strip() for p in target_clean.split(',') if p.strip()]

    if cand_clean == target_clean:
        return {"relevancy": "High", "status": "Exact match (Assumed Local)"}
    if len(cand_parts) == 1 and cand_parts[0] in target_parts:
        return {"relevancy": "High", "status": f"Broad city match - {cand_parts[0].title()} (Assumed Local)"}
    if len(target_parts) == 1 and target_parts[0] in cand_parts:
        return {"relevancy": "High", "status": f"Broad city match - {target_parts[0].title()} (Assumed Local)"}

    cand_data = get_coordinates(candidate_city)
    office_data = get_coordinates(target_city)

    if not cand_data or not office_data:
        return {"relevancy": "Unknown", "status": "Location unknown."}

    cand_coords = (cand_data[0], cand_data[1])
    office_coords = (office_data[0], office_data[1])
    cand_address = cand_data[2].lower() 

    cand_rad = np.radians([cand_coords])
    office_rad = np.radians([office_coords])
    
    dist_matrix = haversine_distances(office_rad, cand_rad)
    straight_line_km = dist_matrix[0][0] * 6371.0  
    
    if straight_line_km > 50.0:
        if "maharashtra" in cand_address or "maharashtra" in candidate_city.lower():
            return {"relevancy": "Low", "status": f"In-State ({straight_line_km:.1f} km away)"}
        else:
            return {"relevancy": "Relocation", "status": f"Out of State ({straight_line_km:.1f} km away)"}

    lon1, lat1 = office_coords[1], office_coords[0]
    lon2, lat2 = cand_coords[1], cand_coords[0]
    
    try:
        osrm_url = f"http://router.project-osrm.org/route/v1/driving/{lon1},{lat1};{lon2},{lat2}?overview=false"
        response = requests.get(osrm_url, timeout=5)
        data = response.json()
        
        if data.get("code") == "Ok":
            duration_seconds = data["routes"][0]["duration"]
            duration_mins = int(duration_seconds // 60)
            
            if duration_mins < 45:
                return {"relevancy": "High", "status": f"~{duration_mins} min drive (Excellent Commute)"}
            elif duration_mins <= 90:
                return {"relevancy": "Medium", "status": f"~{duration_mins} min drive (Moderate Commute)"}
            else:
                return {"relevancy": "Low", "status": f"~{duration_mins} min drive (Tough Commute)"}
    except Exception as e:
        logger.error(f"OSRM Routing failed: {e}")
    
    return {"relevancy": "Unknown", "status": f"{straight_line_km:.1f} km (API limits reached)"}

# --- 2. FASTAPI SETUP ---
@asynccontextmanager
async def lifespan(app: FastAPI):
    yield

app = FastAPI(title="Muuchstac ATS Engine", lifespan=lifespan)
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_methods=["*"], allow_headers=["*"])

@dataclass
class FilterConfig:
    min_exp: float
    required_skills: str
    required_education: str
    target_loc: str
    mand_exp: bool
    mand_edu: bool
    mand_skill: bool
    mand_loc: bool
    passing_score: int

# 👉 NEW: Schema for JD Extraction
class JDExtractionAI(BaseModel):
    min_experience_years: float = Field(description="Minimum years of experience required (output 0 if not mentioned)")
    required_skills: str = Field(description="Top 5-7 core technical skills required, separated by commas")
    required_education: str = Field(description="The minimum education degree required (e.g., 'Bachelors', 'MBA'). Leave blank if not mentioned.")
    target_location: str = Field(description="The primary city or location for the job. Leave blank if remote/not mentioned.")

class CandidateEvaluationAI(BaseModel):
    candidate_name: str       = Field(description="Full name extracted from resume")
    experience_score: int     = Field(description="Experience score out of 40")
    experience_details: str   = Field(description="1 short sentence explaining why they got this experience score")
    skills_score: int         = Field(description="Skills score out of 30")
    skills_details: str       = Field(description="1 short sentence explaining why they got this skills score")
    education_score: int      = Field(description="Education score out of 30")
    education_details: str    = Field(description="1 short sentence explaining why they got this education score")
    score_justification: str  = Field(description="One sentence overall justification")
    candidate_location: str   = Field(description="Specific City or area where candidate lives (e.g., 'Navi Mumbai', 'Pune')")
    contact_email: str        = Field(description="Candidate email")
    contact_phone: str        = Field(description="Candidate phone")
    experience_years: float   = Field(description="Years of relevant experience")
    skills: list[str]         = Field(description="Top matching skills")
    missing_requirements: list[str] = Field(description="Missing requirements")

class CandidateEvaluation(CandidateEvaluationAI):
    total_score: int = 0
    location_relevancy: str = ""
    location_details: str = ""
    is_qualified: bool = True
    source_file: str = Field(default="")

# --- 3. HELPER FUNCTIONS ---
def extract_text_from_bytes(file_bytes: bytes, filename: str) -> str:
    try:
        if filename.lower().endswith(".pdf"):
            reader = pypdf.PdfReader(io.BytesIO(file_bytes))
            return " ".join(p.extract_text() for p in reader.pages if p.extract_text())
        elif filename.lower().endswith(".docx"):
            doc = docx.Document(io.BytesIO(file_bytes))
            parts = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text: parts.append(row_text)
            return " ".join(parts)
        raise ValueError(f"Unsupported file type")
    except Exception as e:
        raise ValueError(f"Could not read: {str(e)}")

def build_prompt(resume_text: str, jd: str, cfg: FilterConfig) -> str:
    return f"""You are an HR Gatekeeper. Respond in JSON. 

JD: {jd}
RESUME: {resume_text}

REQUIREMENTS:
Min Exp: {cfg.min_exp} yrs (Mandatory: {cfg.mand_exp})

SCORING (100 pts total for AI):
Exp (40 pts)
Skills (30 pts)
Edu (30 pts)

RULES:
1. If a requirement is Mandatory=True and candidate fails it, set is_qualified = false.
2. DO NOT score Location. Just extract the precise city/neighborhood for candidate_location.
3. Carefully calculate the candidate's total years of professional work experience and output it as a number in experience_years (e.g., 2.5 for two and a half years).
"""

# 👉 NEW: The AI Router Function
async def call_ai_engine(prompt: str, ai_provider: str) -> dict:
    if ai_provider == "gemini":
        response = await asyncio.to_thread(
            client.models.generate_content,
            model=GEMINI_MODEL,
            contents=prompt,
            config=types.GenerateContentConfig(response_mime_type="application/json", response_schema=CandidateEvaluationAI)
        )
        return json.loads(response.text)
        
    elif ai_provider == "openai":
        # Placeholder for future OpenAI integration
        raise NotImplementedError("OpenAI API is not yet configured.")
        
    elif ai_provider == "custom":
        # Placeholder for future Custom / Local API integration
        raise NotImplementedError("Custom API is not yet configured.")
        
    else:
        raise ValueError(f"Unknown AI Provider: {ai_provider}")

async def evaluate_resume(file_bytes: bytes, filename: str, jd: str, cfg: FilterConfig, ai_provider: str) -> dict:
    async with semaphore:
        try:
            text = await asyncio.to_thread(extract_text_from_bytes, file_bytes, filename)
            
            if not text or len(text.strip()) < 20:
                logger.error(f"Failed '{filename}': No readable text found.")
                return {"_failed": True, "_reason": f"Failed '{filename}': Could not read text (Scanned image or empty)."}

            prompt = build_prompt(text, jd, cfg)
            
            # 👉 Pass the prompt to the new router
            result = await call_ai_engine(prompt, ai_provider)
            result["is_qualified"] = True 
            
            geo_data = await asyncio.to_thread(calculate_hybrid_location_score, result["candidate_location"], cfg.target_loc)
            
            result["location_relevancy"] = geo_data["relevancy"]
            result["location_details"] = geo_data["status"]
            result["candidate_location"] = f'{result["candidate_location"]} - {geo_data["status"]}'

            if cfg.mand_loc and result["location_relevancy"] == "Relocation":
                result["is_qualified"] = False
                result["score_justification"] = f"Rejected: Location commute too far ({geo_data['status']})."

            calculated_total = result.get("experience_score", 0) + result.get("skills_score", 0) + result.get("education_score", 0)
            result["total_score"] = calculated_total

            if calculated_total < cfg.passing_score:
                result["is_qualified"] = False
                if "Rejected:" not in result.get("score_justification", ""):
                    result["score_justification"] = f"Rejected: Total score {calculated_total}/100 is below passing threshold."

            result["source_file"] = filename          
            return result
            
        except Exception as exc:
            logger.error(f"Pipeline crashed for {filename}: {exc}")
            return {"_failed": True, "_reason": f"Failed '{filename}': {exc}"}

# --- 4. API ENDPOINT ---

# 👉 NEW: The Endpoint that reads the JD
@app.post("/extract-jd-params/")
async def extract_jd_params(job_description: str = Form(...), ai_provider: str = Form("gemini")):
    prompt = f"""
    You are an expert technical recruiter. Read the following Job Description and extract the core requirements.
    Return the data exactly as requested in the JSON schema.
    
    JOB DESCRIPTION:
    {job_description}
    """
    
    try:
        # We reuse your existing AI Router to do the extraction!
        if ai_provider == "gemini":
            response = await asyncio.to_thread(
                client.models.generate_content,
                model=GEMINI_MODEL,
                contents=prompt,
                config=types.GenerateContentConfig(response_mime_type="application/json", response_schema=JDExtractionAI)
            )
            return json.loads(response.text)
        else:
            raise NotImplementedError("Only Gemini is configured for JD extraction right now.")
    except Exception as e:
        logger.error(f"JD Extraction failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/analyze-batch-parallel/", response_model=list[CandidateEvaluation])
async def analyze_batch_parallel(
    response: Response,   
    files: List[UploadFile] = File(...),
    job_description: str = Form(...),
    min_experience_years: float = Form(0.0),
    mandatory_experience: bool = Form(False),
    required_skills: str = Form(""),
    mandatory_skills: bool = Form(False),
    required_education: str = Form(""),
    mandatory_education: bool = Form(False),
    target_location: str = Form("Borivali, Mumbai"),
    mandatory_location: bool = Form(False),
    passing_score: int = Form(60),
    shortlist_top_n: int = Form(0),
    ai_provider: str = Form("gemini") # 👉 NEW: Receive the engine choice
):
    # 👉 TIER 1 DEDUPLICATION: Byte-Level Hashing
    file_data = []
    seen_hashes = set()
    
    for f in files:
        if f.filename.lower().endswith((".pdf", ".docx")):
            content = await f.read()
            file_hash = hashlib.md5(content).hexdigest() 
            
            if file_hash not in seen_hashes:
                seen_hashes.add(file_hash)
                file_data.append((content, f.filename))
            else:
                logger.info(f"Tier 1 Duplicate Blocked: Skipped '{f.filename}' (Identical File)")

    if not file_data: 
        raise HTTPException(status_code=400, detail="No valid or unique files found.")

    cfg = FilterConfig(min_experience_years, required_skills, required_education, target_location, 
                       mandatory_experience, mandatory_education, mandatory_skills, mandatory_location, passing_score)

    # Pass the ai_provider into the processing loop
    tasks = [evaluate_resume(fb, fn, job_description, cfg, ai_provider) for fb, fn in file_data]
    raw_results = await asyncio.gather(*tasks)

    successes_raw = [r for r in raw_results if not r.get("_failed")]
    if not successes_raw: 
        raise HTTPException(status_code=500, detail="All files failed (Check terminal for details).")

    # 👉 TIER 2 DEDUPLICATION: Identity Matching
    unique_candidates = {}
    
    for cand in successes_raw:
        email = str(cand.get("contact_email", "")).lower().strip()
        name = str(cand.get("candidate_name", "")).lower().strip()
        
        if email and email not in ["not found", "n/a", "unknown", "none", ""]:
            cand_id = email
        elif name and name not in ["not found", "n/a", "unknown", "none", ""]:
            cand_id = name
        else:
            cand_id = str(id(cand)) 
            
        if cand_id not in unique_candidates:
            unique_candidates[cand_id] = cand
        else:
            existing_score = unique_candidates[cand_id].get("total_score", 0)
            new_score = cand.get("total_score", 0)
            
            if new_score > existing_score:
                unique_candidates[cand_id] = cand
                logger.info(f"Tier 2 Duplicate Merged: Kept higher scoring resume for {cand_id}")
            else:
                logger.info(f"Tier 2 Duplicate Dropped: Ignored lower/equal scoring resume for {cand_id}")

    successes = list(unique_candidates.values())
    
    successes.sort(key=lambda c: (not c.get("is_qualified", False), -c.get("total_score", 0)))
    return successes[:shortlist_top_n] if shortlist_top_n > 0 else successes